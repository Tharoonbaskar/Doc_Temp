from docx import Document

d = Document()

d.add_heading('SANCTION LETTER', 0)

d.add_paragraph(
    "Date: <DATE>\n"
    "Reference No.: <TRN>\n"
    "Branch: <BRANCH_NAME>\n"
    "Branch Email: <BRANCH_EMAIL>"
)

d.add_paragraph("To,\n<CUSTOMER_NAME>\n<ADDRESS_TABLE>")

d.add_paragraph("Subject: Sanction of Loan - <LOAN_ACCOUNT_NUMBER>")

d.add_paragraph("Dear Sir/Madam,")

d.add_paragraph(
    "We are pleased to inform you that your request for financial assistance "
    "has been approved by the Bank, subject to terms and conditions."
)

d.add_paragraph("Customer ID: <CUSTOMER_ID>")
d.add_paragraph("Sanctioned Amount: <LOAN_AMOUNT>")
d.add_paragraph("Interest Rate (p.a.): <ROI>")
d.add_paragraph("EMI Amount: <EMI>")

d.add_paragraph("<PAYMENT_SCHEDULE>")
d.add_paragraph("<CO_APPLICANT_TABLE>")

d.add_paragraph("Borrower Signature: <SIGNATURE>")
d.add_paragraph("Co-applicant Signature: <CO_APPLICANT_SIGNATURE>")
d.add_paragraph("Authorised Signatory: <AUTHORIZED_SIGNATORY>")

d.add_paragraph("<SIGNATURE_TABLE>")

d.save("apps/templates/sample_sanction_letter.docx")

print("Document created successfully")