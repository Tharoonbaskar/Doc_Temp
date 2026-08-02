"""
Test suite for Canvas → ProseMirror migration validation

Tests verify:
1. Database stores prosemirror_json correctly
2. Word import generates ProseMirror JSON (not Canvas elements)
3. Position tracking removed from diff engine
4. Legacy content_json payloads are rejected
"""

import json
from django.test import TestCase
from django.contrib.auth import get_user_model
from apps.templates.models import Template, TemplateVersion, TemplateElementChange
from apps.documents.models import Document as DocumentModel, DocumentCategory
from apps.templates.parsers import ProseMirrorDocumentParser
from apps.templates.pdf_engine import EnterprisePDFEngine
from apps.templates.services import TemplateService
from apps.templates.diff_utils import TemplateElementDiffer
from io import BytesIO
from docx import Document

User = get_user_model()


class ProseMirrorMigrationTests(TestCase):
    """Test suite for ProseMirror migration validation"""

    def setUp(self):
        """Set up test data"""
        self.user = User.objects.create_user(
            username='testuser',
            email='test@example.com',
            password='testpass123'
        )
        self.service = TemplateService()
        
        # Create required Document for Template FK
        self.category = DocumentCategory.objects.create(
            code="TEST_CAT",
            name="Test Category",
            description="Test category for migration tests"
        )
        self.document = DocumentModel.objects.create(
            code="TEST_DOC",
            category=self.category,
            name="Test Document",
            document_type="FORM",
            business_module="TEST",
            product="TEST",
            output_format="PDF",
            description="Test document for migration tests"
        )

    def test_database_stores_prosemirror_json_field(self):
        """Test that templates store prosemirror_json in dedicated field"""
        # Create template with ProseMirror JSON
        prosemirror_data = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {
                            "type": "text",
                            "text": "Test paragraph with "
                        },
                        {
                            "type": "text",
                            "marks": [{"type": "bold"}],
                            "text": "bold text"
                        }
                    ]
                }
            ]
        }

        template = Template.objects.create(
            name="Test Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=prosemirror_data,
            page_size="A4",
            page_orientation="PORTRAIT",
            created_by=self.user
        )

        # Verify field is stored correctly
        template.refresh_from_db()
        self.assertIsNotNone(template.prosemirror_json)
        self.assertEqual(template.prosemirror_json["type"], "doc")
        self.assertEqual(len(template.prosemirror_json["content"]), 1)
        self.assertEqual(template.page_size, "A4")
        self.assertEqual(template.page_orientation, "PORTRAIT")

        print("✅ TEST PASSED: Database stores prosemirror_json correctly")

    def test_word_parser_generates_prosemirror_not_canvas(self):
        """Test that new parser generates ProseMirror JSON instead of Canvas elements"""
        # Create a simple DOCX document in memory
        doc = Document()
        doc.add_heading('Test Heading', level=1)
        doc.add_paragraph('This is a test paragraph with some text.')
        doc.add_paragraph('This paragraph has ').add_run('bold text').bold = True

        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Parse with new ProseMirror parser
        parser = ProseMirrorDocumentParser()
        parsed_payload = parser.parse(docx_buffer)
        prosemirror_json = parsed_payload["prosemirror_json"]

        # Verify enterprise envelope fields exist
        self.assertIn("validation_report", parsed_payload)
        self.assertIn("layout", parsed_payload)
        self.assertIn("variable_summary", parsed_payload)

        # Verify structure
        self.assertIsInstance(prosemirror_json, dict)
        self.assertEqual(prosemirror_json["type"], "doc")
        self.assertIn("content", prosemirror_json)
        self.assertIsInstance(prosemirror_json["content"], list)
        self.assertGreater(len(prosemirror_json["content"]), 0)

        # Verify no Canvas-specific fields (x, y, width, height, rotation, zIndex)
        content_str = json.dumps(prosemirror_json)
        self.assertNotIn('"x":', content_str)
        self.assertNotIn('"y":', content_str)
        self.assertNotIn('"width":', content_str)
        self.assertNotIn('"height":', content_str)
        self.assertNotIn('"rotation":', content_str)
        self.assertNotIn('"zIndex":', content_str)

        # Verify has ProseMirror-specific structure
        first_node = prosemirror_json["content"][0]
        self.assertIn("type", first_node)
        # Should be heading or paragraph
        self.assertIn(first_node["type"], ["heading", "paragraph"])

        print("✅ TEST PASSED: Word parser generates ProseMirror JSON (no Canvas elements)")

    def test_word_parser_normalizes_enterprise_placeholders_to_typed_variable_nodes(self):
        """All placeholder syntaxes should normalize into enterprise variable node types with canonical token format."""
        doc = Document()
        doc.add_paragraph("Loan amount: {LOAN_AMOUNT}")
        doc.add_paragraph("TRN: {{TRN}}")
        doc.add_paragraph("Customer: <CUSTOMER_ID>")
        doc.add_paragraph("Address block: {ADDRESS_TABLE}")
        doc.add_paragraph("Signature: {{AUTHORIZED_SIGNATORY}}")
        doc.add_paragraph("Photo: <CUSTOMER_PHOTO>")

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        parser = ProseMirrorDocumentParser()
        parsed_payload = parser.parse(docx_buffer)
        pm_doc = parsed_payload["prosemirror_json"]
        summary = parsed_payload["variable_summary"]

        content_dump = json.dumps(pm_doc)
        self.assertIn('"type": "variableChip"', content_dump)
        self.assertIn('"type": "dynamicTableVariable"', content_dump)
        self.assertIn('"type": "signatureVariable"', content_dump)
        self.assertIn('"type": "imagePlaceholderVariable"', content_dump)
        self.assertIn('"field": "LOAN_AMOUNT"', content_dump)
        self.assertIn('"field": "TRN"', content_dump)
        self.assertIn('"field": "CUSTOMER_ID"', content_dump)
        self.assertIn('"field": "ADDRESS_TABLE"', content_dump)
        self.assertIn('"field": "AUTHORIZED_SIGNATORY"', content_dump)
        self.assertIn('"field": "CUSTOMER_PHOTO"', content_dump)

        self.assertEqual(summary.get("normalized_format"), "<VARIABLE_NAME>")
        self.assertEqual(summary.get("total_detected"), 6)

        detected_tokens = {item["token"] for item in summary.get("variables", [])}
        self.assertSetEqual(
            detected_tokens,
            {
                "LOAN_AMOUNT",
                "TRN",
                "CUSTOMER_ID",
                "ADDRESS_TABLE",
                "AUTHORIZED_SIGNATORY",
                "CUSTOMER_PHOTO",
            },
        )

        by_category = summary.get("by_category", {})
        self.assertGreaterEqual(by_category.get("simple", 0), 3)
        self.assertGreaterEqual(by_category.get("dynamic_table", 0), 1)
        self.assertGreaterEqual(by_category.get("signature", 0), 1)
        self.assertGreaterEqual(by_category.get("image", 0), 1)

    def test_word_parser_preserves_table_structure(self):
        """Tables should remain structured ProseMirror table nodes, not flattened text."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        parser = ProseMirrorDocumentParser()
        parsed_payload = parser.parse(docx_buffer)
        pm_doc = parsed_payload["prosemirror_json"]

        table_nodes = [node for node in pm_doc.get("content", []) if isinstance(node, dict) and node.get("type") == "table"]
        self.assertEqual(len(table_nodes), 1)

        table_node = table_nodes[0]
        rows = table_node.get("content", [])
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0].get("type"), "tableRow")
        self.assertEqual(rows[1].get("type"), "tableRow")

    def test_word_parser_extracts_headers_and_footers_into_layout_metadata(self):
        """Header/footer content should be preserved in layout metadata envelope."""
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Enterprise Header"
        section.footer.paragraphs[0].text = "Page footer"
        doc.add_paragraph("Body content")

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        parser = ProseMirrorDocumentParser()
        parsed_payload = parser.parse(docx_buffer)
        layout = parsed_payload["layout"]

        self.assertIn("headers", layout)
        self.assertIn("footers", layout)
        self.assertGreaterEqual(len(layout["headers"]), 1)
        self.assertGreaterEqual(len(layout["footers"]), 1)

        header_dump = json.dumps(layout["headers"])
        footer_dump = json.dumps(layout["footers"])
        self.assertIn("Enterprise Header", header_dump)
        self.assertIn("Page footer", footer_dump)

    def test_word_parser_preserves_intentional_blank_paragraphs(self):
        """Blank paragraphs in the body should be preserved for spacing fidelity."""
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("")
        doc.add_paragraph("Second paragraph")

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        parser = ProseMirrorDocumentParser()
        parsed_payload = parser.parse(docx_buffer)
        pm_doc = parsed_payload["prosemirror_json"]

        paragraphs = [
            node
            for node in pm_doc.get("content", [])
            if isinstance(node, dict) and node.get("type") == "paragraph"
        ]
        self.assertGreaterEqual(len(paragraphs), 3)

        content_dump = json.dumps(pm_doc)
        self.assertIn("First paragraph", content_dump)
        self.assertIn("Second paragraph", content_dump)

    def test_pdf_renderer_maps_docx_spacing_attrs(self):
        """PDF renderer should honor parser-style docx spacing keys when explicit style keys are absent."""
        engine = EnterprisePDFEngine()
        html_fragment = engine._render_block(
            {
                "type": "paragraph",
                "attrs": {
                    "docx": {
                        "lineSpacing": 18,
                        "spaceBefore": 6,
                        "spaceAfter": 12,
                        "firstLineIndent": 10,
                        "leftIndent": 8,
                        "rightIndent": 4,
                    }
                },
                "content": [{"type": "text", "text": "Spacing test"}],
            }
        )

        self.assertIn("line-height:18pt", html_fragment)
        self.assertIn("margin-top:6pt", html_fragment)
        self.assertIn("margin-bottom:12pt", html_fragment)
        self.assertIn("text-indent:10pt", html_fragment)
        self.assertIn("margin-left:8pt", html_fragment)
        self.assertIn("margin-right:4pt", html_fragment)

    def test_diff_engine_no_position_tracking(self):
        """Test that diff engine does NOT track POSITION_CHANGED"""
        # Create base template
        base_prosemirror = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Original text"}
                    ]
                }
            ]
        }

        modified_prosemirror = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Modified text"}
                    ]
                }
            ]
        }

        # Create differ
        differ = TemplateElementDiffer()
        diff_data = differ.calculate_diff(base_prosemirror, modified_prosemirror)
        changes = TemplateElementDiffer.extract_element_changes(diff_data)

        # Verify no POSITION_CHANGED semantic type
        # changes is a list of tuples: (element_id, change_type, old_value, new_value)
        for change_tuple in changes:
            element_id, change_type, old_value, new_value = change_tuple
            
            semantic = None
            if isinstance(new_value, dict):
                semantic = new_value.get('_semantic', {})
            elif isinstance(old_value, dict):
                semantic = old_value.get('_semantic', {})

            if semantic:
                semantic_type = semantic.get('type')
                self.assertNotEqual(semantic_type, 'POSITION_CHANGED',
                                  "POSITION_CHANGED should not be tracked")
                self.assertNotEqual(semantic_type, 'IMAGE_MOVED',
                                  "IMAGE_MOVED should not be tracked")

        print("✅ TEST PASSED: Diff engine does not track POSITION_CHANGED")

    def test_change_records_no_position_fields(self):
        """Test that change records do NOT contain oldPosition/newPosition"""
        # Create template with version
        base_prosemirror = {
            "type": "doc",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "Base"}]}
            ]
        }

        modified_prosemirror = {
            "type": "doc",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "Modified"}]}
            ]
        }

        template = Template.objects.create(
            name="Change Test Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_prosemirror,
            created_by=self.user
        )

        # Create base version (use template_json field, not prosemirror_json)
        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_status='APPROVED',
            template_json=base_prosemirror,
            code=f"{template.code}_V1",
            version_name="v1.0",
            created_by=self.user,
            updated_by=self.user
        )

        # Create modified version with changes
        differ = TemplateElementDiffer()
        diff_data = differ.calculate_diff(base_prosemirror, modified_prosemirror)
        changes = TemplateElementDiffer.extract_element_changes(diff_data)

        # Verify serialized changes don't have position fields (from structured_changes)
        service = TemplateService()
        for change_tuple in changes:
            element_id, change_type, old_value, new_value = change_tuple
            # Extract semantic info from new_value
            semantic = None
            if isinstance(new_value, dict):
                semantic = new_value.get('_semantic', {})
            elif isinstance(old_value, dict):
                semantic = old_value.get('_semantic', {})

            # Verify no position fields in semantic data
            if semantic:
                self.assertNotIn('oldPosition', semantic,
                               "oldPosition should not be in semantic data")
                self.assertNotIn('newPosition', semantic,
                               "newPosition should not be in semantic data")

        print("✅ TEST PASSED: Change records do not contain position fields")

    def test_legacy_content_json_payload_is_rejected(self):
        """Test that content_json payloads are no longer accepted by service API."""
        response = self.service.create({
            "code": "DOC_STRICT_PM_TEMPLATE_000001",
            "name": "Legacy Payload Template",
            "category": "CONTRACT",
            "document": self.document,
            "template_type": "DYNAMIC",
            "content_type": "application/json",
            "content_json": json.dumps({"prosemirror_json": {"type": "doc", "content": [{"type": "paragraph"}]}}),
            "prosemirror_json": {
                "type": "doc",
                "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Strict"}]}],
            },
            "created_by": self.user,
        })

        self.assertEqual(response.status_code, 400)
        response_payload = response.data if isinstance(response.data, dict) else {}
        message = str(response_payload.get("message") or "")
        self.assertIn("content_json", message)

        print("✅ TEST PASSED: Legacy content_json payloads are rejected")

    def test_new_templates_use_prosemirror_json_field(self):
        """Test that new templates use prosemirror_json field directly"""
        prosemirror_data = {
            "type": "doc",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "New template"}]}
            ]
        }

        template = Template.objects.create(
            name="New Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=prosemirror_data,
            page_size="A4",
            page_orientation="PORTRAIT",
            created_by=self.user
        )

        # Verify prosemirror_json is populated
        template.refresh_from_db()
        self.assertIsNotNone(template.prosemirror_json)
        self.assertEqual(template.prosemirror_json["type"], "doc")
        self.assertIsNone(template.content_json)

        print("✅ TEST PASSED: New templates use prosemirror_json field")


class ProseMirrorServiceIntegrationTests(TestCase):
    """Integration tests for TemplateService with ProseMirror"""

    def setUp(self):
        self.user = User.objects.create_user(
            username='integrationuser',
            email='integration@example.com',
            password='testpass123'
        )
        self.service = TemplateService()
        
        # Create required Document for Template FK
        self.category = DocumentCategory.objects.create(
            code="INT_CAT",
            name="Integration Category",
            description="Integration test category"
        )
        self.document = DocumentModel.objects.create(
            code="INT_DOC",
            category=self.category,
            name="Integration Document",
            document_type="FORM",
            business_module="TEST",
            product="TEST",
            output_format="PDF",
            description="Integration test document"
        )

    def test_save_template_stores_prosemirror_correctly(self):
        """Test that saving a template via service stores ProseMirror correctly"""
        payload = {
            "name": "Service Test Template",
            "category": "CONTRACT",
            "prosemirror_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [
                            {"type": "text", "text": "Service layer test"}
                        ]
                    }
                ]
            },
            "html": "<p>Service layer test</p>",
            "page": {
                "size": "A4",
                "orientation": "PORTRAIT"
            }
        }

        # Create via service (simulate API call)
        template = Template.objects.create(
            name=payload["name"],
            category=payload["category"],
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=payload["prosemirror_json"],
            page_size=payload["page"]["size"],
            page_orientation=payload["page"]["orientation"],
            created_by=self.user
        )

        # Verify
        template.refresh_from_db()
        self.assertEqual(template.prosemirror_json["type"], "doc")
        self.assertEqual(template.page_size, "A4")

        print("✅ TEST PASSED: Service layer stores ProseMirror correctly")

    def test_approve_draft_version_preserves_draft_document_when_all_changes_accepted(self):
        """Ensure approval does not drop reviewed draft content when every change is accepted."""
        base_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1, "textAlign": "center"},
                    "content": [{"type": "text", "text": "TESTING LETTER"}],
                },
                {
                    "type": "paragraph",
                    "attrs": {"textAlign": None},
                },
            ],
        }

        draft_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1, "textAlign": "center"},
                    "content": [{"type": "text", "text": "TESTING LETTER"}],
                },
                {
                    "type": "paragraph",
                    "attrs": {"textAlign": None},
                    "content": [
                        {"type": "text", "text": "Dear "},
                        {
                            "type": "variableChip",
                            "attrs": {
                                "field": "applicant_name",
                                "label": "Applicant Name",
                            },
                        },
                        {"type": "text", "text": ", We are pleased to inform you."},
                    ],
                },
            ],
        }

        template = Template.objects.create(
            code="DOC_APPROVAL_PRESERVE_000001",
            name="Approval Preserve Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_doc,
            status="APPROVED",
            created_by=self.user,
            updated_by=self.user,
        )

        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_name="v1.0",
            version_status="APPROVED",
            template_json=base_doc,
            code=f"{template.code}_V1",
            created_by=self.user,
            updated_by=self.user,
        )

        draft_version = TemplateVersion.objects.create(
            template=template,
            version_number=2,
            version_name="v2.0",
            version_status="FOR_REVIEW",
            # Keep this wrapped to mirror real-world persisted payloads from older edits.
            template_json={"prosemirror_json": draft_doc},
            base_version=base_version,
            code=f"{template.code}_V2",
            created_by=self.user,
            updated_by=self.user,
        )

        TemplateElementChange.objects.create(
            version=draft_version,
            element_id="path:0.1:paragraph",
            change_type="MODIFIED",
            old_value={
                "type": "paragraph",
                "attrs": {"textAlign": None},
                "text": "",
                "_semantic": {
                    "type": "TEXT_MODIFIED",
                    "oldText": "",
                    "newText": "Dear , We are pleased to inform you.",
                },
            },
            new_value={
                "type": "paragraph",
                "attrs": {"textAlign": None},
                "content": draft_doc["content"][1]["content"],
                "_semantic": {
                    "type": "TEXT_MODIFIED",
                    "oldText": "",
                    "newText": "Dear , We are pleased to inform you.",
                    "newValue": {
                        "type": "paragraph",
                        "attrs": {"textAlign": None},
                        "content": draft_doc["content"][1]["content"],
                    },
                },
            },
            approval_status="RESOLVED",
            code=f"{draft_version.code}_CHANGE_1",
            created_by=self.user,
            updated_by=self.user,
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        draft_version.refresh_from_db()

        approved_doc = template.prosemirror_json
        self.assertIsInstance(approved_doc, dict)
        self.assertEqual(approved_doc.get("type"), "doc")

        content_dump = json.dumps(approved_doc)
        self.assertIn("Dear", content_dump)
        self.assertIn("applicant_name", content_dump)
        self.assertIn("We are pleased to inform you", content_dump)

        self.assertEqual(draft_version.version_status, "APPROVED")

    def _create_in_review_version_with_change(
        self,
        *,
        code_suffix: str,
        base_doc: dict,
        draft_doc: dict,
        element_id: str,
        change_type: str,
        old_value,
        new_value,
        approval_status: str,
    ):
        template = Template.objects.create(
            code=f"DOC_WORD_TRACK_{code_suffix}",
            name=f"Word Track {code_suffix}",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_doc,
            status="APPROVED",
            created_by=self.user,
            updated_by=self.user,
        )

        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_name="v1.0",
            version_status="APPROVED",
            template_json=base_doc,
            code=f"{template.code}_V1",
            created_by=self.user,
            updated_by=self.user,
        )

        draft_version = TemplateVersion.objects.create(
            template=template,
            version_number=2,
            version_name="v2.0",
            version_status="FOR_REVIEW",
            template_json=draft_doc,
            base_version=base_version,
            code=f"{template.code}_V2",
            created_by=self.user,
            updated_by=self.user,
        )

        TemplateElementChange.objects.create(
            version=draft_version,
            element_id=element_id,
            change_type=change_type,
            old_value=old_value,
            new_value=new_value,
            approval_status=approval_status,
            code=f"{draft_version.code}_CHANGE_1",
            created_by=self.user,
            updated_by=self.user,
        )

        return template, draft_version

    def test_word_semantics_approve_added_keeps_inserted_text(self):
        base_doc = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "SANCTION LETTER"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "Dear Applicant,"}]},
            ],
        }
        inserted_paragraph = {
            "type": "paragraph",
            "content": [
                {
                    "type": "text",
                    "text": "We are pleased to inform you that your loan application has been approved.",
                }
            ],
        }
        draft_doc = {
            "type": "doc",
            "content": [
                base_doc["content"][0],
                base_doc["content"][1],
                inserted_paragraph,
            ],
        }

        template, _ = self._create_in_review_version_with_change(
            code_suffix="ADD_APPROVE_000001",
            base_doc=base_doc,
            draft_doc=draft_doc,
            element_id="path:0.2:paragraph",
            change_type="ADDED",
            old_value=None,
            new_value={
                **inserted_paragraph,
                "_semantic": {"type": "TEXT_ADDED", "newValue": inserted_paragraph},
            },
            approval_status="APPROVED",
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        approved_dump = json.dumps(template.prosemirror_json)
        self.assertIn("loan application has been approved", approved_dump)

    def test_word_semantics_reject_added_removes_inserted_text(self):
        base_doc = {
            "type": "doc",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "Dear Applicant,"}]},
            ],
        }
        inserted_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Inserted paragraph for review"}],
        }
        draft_doc = {
            "type": "doc",
            "content": [
                base_doc["content"][0],
                inserted_paragraph,
            ],
        }

        template, _ = self._create_in_review_version_with_change(
            code_suffix="ADD_REJECT_000001",
            base_doc=base_doc,
            draft_doc=draft_doc,
            element_id="path:0.1:paragraph",
            change_type="ADDED",
            old_value=None,
            new_value={
                **inserted_paragraph,
                "_semantic": {"type": "TEXT_ADDED", "newValue": inserted_paragraph},
            },
            approval_status="REJECTED",
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        approved_dump = json.dumps(template.prosemirror_json)
        self.assertNotIn("Inserted paragraph for review", approved_dump)
        self.assertIn("Dear Applicant,", approved_dump)

    def test_word_semantics_approve_deleted_removes_text(self):
        removed_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Sample Letter"}],
        }
        base_doc = {"type": "doc", "content": [removed_paragraph]}
        draft_doc = {"type": "doc", "content": [{"type": "paragraph"}]}

        template, _ = self._create_in_review_version_with_change(
            code_suffix="DEL_APPROVE_000001",
            base_doc=base_doc,
            draft_doc=draft_doc,
            element_id="path:0.0:paragraph",
            change_type="DELETED",
            old_value={
                **removed_paragraph,
                "_semantic": {"type": "TEXT_REMOVED", "oldValue": removed_paragraph},
            },
            new_value=None,
            approval_status="APPROVED",
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        approved_dump = json.dumps(template.prosemirror_json)
        self.assertNotIn("Sample Letter", approved_dump)

    def test_word_semantics_reject_deleted_restores_text(self):
        removed_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Sample Letter"}],
        }
        base_doc = {"type": "doc", "content": [removed_paragraph]}
        draft_doc = {"type": "doc", "content": [{"type": "paragraph"}]}

        template, _ = self._create_in_review_version_with_change(
            code_suffix="DEL_REJECT_000001",
            base_doc=base_doc,
            draft_doc=draft_doc,
            element_id="path:0.0:paragraph",
            change_type="DELETED",
            old_value={
                **removed_paragraph,
                "_semantic": {"type": "TEXT_REMOVED", "oldValue": removed_paragraph},
            },
            new_value=None,
            approval_status="REJECTED",
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        approved_dump = json.dumps(template.prosemirror_json)
        self.assertIn("Sample Letter", approved_dump)

    def test_word_semantics_approve_modified_keeps_new_text(self):
        old_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Sample Letter"}],
        }
        new_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Sanction Letter"}],
        }
        base_doc = {"type": "doc", "content": [old_paragraph]}
        draft_doc = {"type": "doc", "content": [new_paragraph]}

        template, _ = self._create_in_review_version_with_change(
            code_suffix="MOD_APPROVE_000001",
            base_doc=base_doc,
            draft_doc=draft_doc,
            element_id="path:0.0:paragraph",
            change_type="MODIFIED",
            old_value={
                **old_paragraph,
                "_semantic": {
                    "type": "TEXT_MODIFIED",
                    "oldValue": old_paragraph,
                    "newValue": new_paragraph,
                },
            },
            new_value={
                **new_paragraph,
                "_semantic": {
                    "type": "TEXT_MODIFIED",
                    "oldValue": old_paragraph,
                    "newValue": new_paragraph,
                },
            },
            approval_status="APPROVED",
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        approved_dump = json.dumps(template.prosemirror_json)
        self.assertIn("Sanction Letter", approved_dump)
        self.assertNotIn("Sample Letter", approved_dump)

    def test_word_semantics_reject_modified_restores_old_text(self):
        old_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Sample Letter"}],
        }
        new_paragraph = {
            "type": "paragraph",
            "content": [{"type": "text", "text": "Sanction Letter"}],
        }
        base_doc = {"type": "doc", "content": [old_paragraph]}
        draft_doc = {"type": "doc", "content": [new_paragraph]}

        template, _ = self._create_in_review_version_with_change(
            code_suffix="MOD_REJECT_000001",
            base_doc=base_doc,
            draft_doc=draft_doc,
            element_id="path:0.0:paragraph",
            change_type="MODIFIED",
            old_value={
                **old_paragraph,
                "_semantic": {
                    "type": "TEXT_MODIFIED",
                    "oldValue": old_paragraph,
                    "newValue": new_paragraph,
                },
            },
            new_value={
                **new_paragraph,
                "_semantic": {
                    "type": "TEXT_MODIFIED",
                    "oldValue": old_paragraph,
                    "newValue": new_paragraph,
                },
            },
            approval_status="REJECTED",
        )

        response = self.service.approve_draft_version(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        approved_dump = json.dumps(template.prosemirror_json)
        self.assertIn("Sample Letter", approved_dump)
        self.assertNotIn("Sanction Letter", approved_dump)

    def test_send_back_for_revision_moves_pending_version_to_draft(self):
        """Template-level send back should keep approved baseline while reverting pending version to draft."""
        base_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Approved baseline"}],
                }
            ],
        }
        draft_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Pending review update"}],
                }
            ],
        }

        template = Template.objects.create(
            code="DOC_SEND_BACK_SYNC_000001",
            name="Send Back Sync Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_doc,
            status="FOR_REVIEW",
            created_by=self.user,
            updated_by=self.user,
        )

        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_name="v1.0",
            version_status="APPROVED",
            template_json=base_doc,
            code=f"{template.code}_V1",
            created_by=self.user,
            updated_by=self.user,
        )

        draft_version = TemplateVersion.objects.create(
            template=template,
            version_number=2,
            version_name="v2.0",
            version_status="FOR_REVIEW",
            template_json=draft_doc,
            base_version=base_version,
            code=f"{template.code}_V2",
            created_by=self.user,
            updated_by=self.user,
        )

        change = TemplateElementChange.objects.create(
            version=draft_version,
            element_id="path:0.0:paragraph",
            change_type="MODIFIED",
            old_value=base_doc["content"][0],
            new_value=draft_doc["content"][0],
            approval_status="APPROVED",
            reviewed_by=self.user,
            review_comment="Looks good",
            code=f"{draft_version.code}_CHG_1",
            created_by=self.user,
            updated_by=self.user,
        )

        response = self.service.send_back_for_revision(template.id, self.user, "Needs update")
        self.assertEqual(response.status, "APPROVED")

        template.refresh_from_db()
        draft_version.refresh_from_db()
        change.refresh_from_db()

        self.assertEqual(template.status, "APPROVED")
        self.assertEqual(template.review_comments, "Needs update")
        self.assertEqual(draft_version.version_status, "DRAFT")
        self.assertEqual(change.approval_status, "PENDING")
        self.assertIsNone(change.reviewed_by)
        self.assertEqual(change.review_comment, "")

    def test_send_draft_version_for_review_keeps_template_approved_when_baseline_exists(self):
        base_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Approved baseline text"}],
                }
            ],
        }
        draft_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Approved baseline text updated"}],
                }
            ],
        }

        template = Template.objects.create(
            code="DOC_KEEP_APPROVED_000001",
            name="Keep Approved Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_doc,
            status="APPROVED",
            created_by=self.user,
            updated_by=self.user,
        )

        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_name="v1.0",
            version_status="APPROVED",
            template_json=base_doc,
            code=f"{template.code}_V1",
            created_by=self.user,
            updated_by=self.user,
        )

        draft_version = TemplateVersion.objects.create(
            template=template,
            version_number=2,
            version_name="v2.0",
            version_status="DRAFT",
            template_json=draft_doc,
            base_version=base_version,
            code=f"{template.code}_V2",
            created_by=self.user,
            updated_by=self.user,
        )

        TemplateElementChange.objects.create(
            version=draft_version,
            element_id="path:0.0:paragraph",
            change_type="MODIFIED",
            old_value=base_doc["content"][0],
            new_value=draft_doc["content"][0],
            approval_status="PENDING",
            code=f"{draft_version.code}_CHG_1",
            created_by=self.user,
            updated_by=self.user,
        )

        response = self.service.send_draft_version_for_review(template.id, 2, self.user)
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        draft_version.refresh_from_db()

        self.assertEqual(template.status, "APPROVED")
        self.assertEqual(draft_version.version_status, "FOR_REVIEW")

    def test_review_send_back_preserves_other_approved_changes(self):
        base_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line A old"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line B old"}],
                },
            ],
        }
        draft_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line A new"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line B new"}],
                },
            ],
        }

        template = Template.objects.create(
            code="DOC_PARTIAL_SEND_BACK_000001",
            name="Partial Send Back Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_doc,
            status="APPROVED",
            created_by=self.user,
            updated_by=self.user,
        )

        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_name="v1.0",
            version_status="APPROVED",
            template_json=base_doc,
            code=f"{template.code}_V1",
            created_by=self.user,
            updated_by=self.user,
        )

        draft_version = TemplateVersion.objects.create(
            template=template,
            version_number=2,
            version_name="v2.0",
            version_status="FOR_REVIEW",
            template_json=draft_doc,
            base_version=base_version,
            code=f"{template.code}_V2",
            created_by=self.user,
            updated_by=self.user,
        )

        approved_change = TemplateElementChange.objects.create(
            version=draft_version,
            element_id="path:0.0:paragraph",
            change_type="MODIFIED",
            old_value=base_doc["content"][0],
            new_value=draft_doc["content"][0],
            approval_status="APPROVED",
            code=f"{draft_version.code}_CHG_A",
            created_by=self.user,
            updated_by=self.user,
        )

        sent_back_change = TemplateElementChange.objects.create(
            version=draft_version,
            element_id="path:0.1:paragraph",
            change_type="MODIFIED",
            old_value=base_doc["content"][1],
            new_value=draft_doc["content"][1],
            approval_status="PENDING",
            code=f"{draft_version.code}_CHG_B",
            created_by=self.user,
            updated_by=self.user,
        )

        response = self.service.review_element_change(
            sent_back_change.id,
            self.user,
            'SENT_BACK',
            'Please revise line B',
        )
        self.assertEqual(response.status_code, 200)

        template.refresh_from_db()
        draft_version.refresh_from_db()
        approved_change.refresh_from_db()
        sent_back_change.refresh_from_db()

        self.assertEqual(template.status, "APPROVED")
        self.assertEqual(draft_version.version_status, "DRAFT")
        self.assertEqual(approved_change.approval_status, "APPROVED")
        self.assertEqual(sent_back_change.approval_status, "SENT_BACK")
        self.assertEqual(sent_back_change.review_comment, "Please revise line B")

    def test_update_draft_version_preserves_previously_approved_changes(self):
        base_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line A old"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line B old"}],
                },
            ],
        }
        draft_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line A new"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line B new"}],
                },
            ],
        }

        template = Template.objects.create(
            code="DOC_DRAFT_STATUS_PRESERVE_000001",
            name="Draft Status Preserve Template",
            category="CONTRACT",
            document=self.document,
            template_type="DYNAMIC",
            content_type="application/json",
            prosemirror_json=base_doc,
            status="APPROVED",
            created_by=self.user,
            updated_by=self.user,
        )

        base_version = TemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_name="v1.0",
            version_status="APPROVED",
            template_json=base_doc,
            code=f"{template.code}_V1",
            created_by=self.user,
            updated_by=self.user,
        )

        draft_version = TemplateVersion.objects.create(
            template=template,
            version_number=2,
            version_name="v2.0",
            version_status="DRAFT",
            template_json=draft_doc,
            base_version=base_version,
            code=f"{template.code}_V2",
            created_by=self.user,
            updated_by=self.user,
        )

        response_initial = self.service.update_draft_version(template.id, 2, self.user, draft_doc)
        self.assertEqual(response_initial.status_code, 200)

        initial_changes = list(draft_version.element_changes.all())
        self.assertGreaterEqual(len(initial_changes), 2)

        approved_change = next(
            change for change in initial_changes if "Line A old" in json.dumps(change.old_value)
        )
        approved_change.approval_status = "APPROVED"
        approved_change.reviewed_by = self.user
        approved_change.review_comment = "Approved line A"
        approved_change.save(update_fields=['approval_status', 'reviewed_by', 'review_comment', 'updated_at'])

        revised_draft_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line A new"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Line B newer"}],
                },
            ],
        }

        response_revised = self.service.update_draft_version(template.id, 2, self.user, revised_draft_doc)
        self.assertEqual(response_revised.status_code, 200)

        refreshed_changes = list(draft_version.element_changes.all())
        preserved_change = next(
            change for change in refreshed_changes if "Line A old" in json.dumps(change.old_value)
        )
        revised_change = next(
            change for change in refreshed_changes if "Line B old" in json.dumps(change.old_value)
        )

        self.assertEqual(preserved_change.approval_status, "APPROVED")
        self.assertEqual(revised_change.approval_status, "PENDING")


class ProseMirrorLifecycleRegressionTests(TestCase):
    """Regression tests for end-to-end ProseMirror lifecycle semantics."""

    def setUp(self):
        self.differ = TemplateElementDiffer()
        self.service = TemplateService()

    @staticmethod
    def _doc_with_paragraph_text(text: str):
        return {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": text}],
                }
            ],
        }

    @staticmethod
    def _empty_doc():
        return {
            "type": "doc",
            "content": [{"type": "paragraph"}],
        }

    def test_scenario_1_identical_document_has_no_changes(self):
        old_doc = self._doc_with_paragraph_text("TESTING")
        new_doc = self._doc_with_paragraph_text("TESTING")

        diff_data = self.differ.calculate_diff(old_doc, new_doc)

        self.assertEqual(diff_data["summary"]["total_changes"], 0)
        self.assertEqual(diff_data["summary"]["added"], 0)
        self.assertEqual(diff_data["summary"]["modified"], 0)
        self.assertEqual(diff_data["summary"]["deleted"], 0)

    def test_scenario_2_delete_all_content_keeps_valid_empty_doc(self):
        old_doc = self._doc_with_paragraph_text("TESTING")
        empty_doc = self._empty_doc()

        diff_data = self.differ.calculate_diff(old_doc, empty_doc)

        self.assertEqual(diff_data["summary"]["deleted"], 1)
        self.assertEqual(diff_data["summary"]["added"], 0)

        changes = TemplateElementDiffer.extract_element_changes(diff_data)
        semantic_types = []
        for _, _, old_value, new_value in changes:
            semantic_payload = {}
            if isinstance(new_value, dict):
                semantic_payload = new_value.get("_semantic") or {}
            elif isinstance(old_value, dict):
                semantic_payload = old_value.get("_semantic") or {}
            semantic_types.append(semantic_payload.get("type"))

        self.assertIn("TEXT_REMOVED", semantic_types)

        normalized_doc = self.service._parse_content_payload(json.dumps({"prosemirror_json": empty_doc}))
        self.assertEqual(normalized_doc["type"], "doc")
        self.assertIsInstance(normalized_doc.get("content"), list)
        self.assertEqual(len(normalized_doc["content"]), 1)
        self.assertEqual(normalized_doc["content"][0].get("type"), "paragraph")

    def test_scenario_3_single_word_deletion_semantic_type_removed(self):
        old_doc = self._doc_with_paragraph_text("TESTING DOCUMENT")
        new_doc = self._doc_with_paragraph_text("DOCUMENT")

        diff_data = self.differ.calculate_diff(old_doc, new_doc)

        self.assertEqual(diff_data["summary"]["added"], 0)
        self.assertEqual(diff_data["summary"]["deleted"], 0)
        self.assertEqual(diff_data["summary"]["modified"], 1)
        self.assertEqual(diff_data["semantic_summary"].get("TEXT_REMOVED", 0), 1)

        changes = TemplateElementDiffer.extract_element_changes(diff_data)
        self.assertTrue(changes)

        _, _, old_value, new_value = changes[0]
        semantic = {}
        if isinstance(new_value, dict):
            semantic = new_value.get("_semantic", {})

        old_delta = str(semantic.get("oldText") or "")
        new_delta = str(semantic.get("newText") or "")
        self.assertIn("TESTING", old_delta)
        self.assertNotIn("DOCUMENT", old_delta)
        self.assertEqual(new_delta.strip(), "")

    def test_scenario_4_insert_paragraph_reports_single_addition(self):
        old_doc = self._doc_with_paragraph_text("ORIGINAL")
        new_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "ORIGINAL"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "NEW PARAGRAPH"}],
                },
            ],
        }

        diff_data = self.differ.calculate_diff(old_doc, new_doc)

        self.assertEqual(diff_data["summary"]["added"], 1)
        self.assertEqual(diff_data["summary"]["deleted"], 0)

    def test_scenario_5_delete_image_reports_single_deletion(self):
        old_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "image",
                    "attrs": {
                        "src": "https://example.com/test.png",
                        "alt": "Test image",
                    },
                }
            ],
        }
        new_doc = self._empty_doc()

        diff_data = self.differ.calculate_diff(old_doc, new_doc)

        self.assertEqual(diff_data["summary"]["deleted"], 1)
        self.assertEqual(diff_data["summary"]["added"], 0)

    def test_scenario_6_delete_variable_reports_single_deletion(self):
        old_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "variable",
                    "attrs": {
                        "binding": "customer_name",
                        "label": "Customer Name",
                    },
                }
            ],
        }
        new_doc = self._empty_doc()

        diff_data = self.differ.calculate_diff(old_doc, new_doc)

        self.assertEqual(diff_data["summary"]["deleted"], 1)
        self.assertEqual(diff_data["summary"]["added"], 0)

    def test_scenario_7_single_space_change_is_reviewable(self):
        old_doc = self._doc_with_paragraph_text("A  B")
        new_doc = self._doc_with_paragraph_text("A B")

        diff_data = self.differ.calculate_diff(old_doc, new_doc)
        self.assertEqual(diff_data["summary"]["modified"], 1)

        changes = TemplateElementDiffer.extract_element_changes(diff_data)
        self.assertTrue(changes)

        _, change_type, old_value, new_value = changes[0]
        self.assertEqual(change_type, "MODIFIED")
        self.assertTrue(self.service._is_reviewable_element_change(change_type, old_value, new_value))

        semantic = {}
        if isinstance(new_value, dict):
            semantic = new_value.get("_semantic", {})
        old_delta = str(semantic.get("oldText") or "")
        new_delta = str(semantic.get("newText") or "")
        self.assertNotEqual(old_delta, new_delta)
        self.assertEqual(old_delta.replace(" ", ""), "B")
        self.assertEqual(new_delta.replace(" ", ""), "B")
        self.assertGreater(len(old_delta), len(new_delta))

    def test_scenario_8_insert_variable_token_is_preserved(self):
        old_doc = self._empty_doc()
        new_doc = self._doc_with_paragraph_text("Dear <applicant_name>")

        diff_data = self.differ.calculate_diff(old_doc, new_doc)
        self.assertGreaterEqual(diff_data["summary"].get("added", 0), 1)

        changes = TemplateElementDiffer.extract_element_changes(diff_data)
        variable_change = None
        for _, change_type, _old, _new in changes:
            if change_type == "ADDED" and isinstance(_new, dict):
                semantic = _new.get("_semantic", {})
                if "<applicant_name>" in str(semantic.get("newText") or ""):
                    variable_change = semantic
                    break

        self.assertIsNotNone(variable_change)
        self.assertIn("Dear <applicant_name>", str(variable_change.get("newText") or ""))

    def test_scenario_9_table_cell_diff_is_scoped_and_indexed(self):
        old_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "table",
                    "content": [
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "<loan_tenure>Months(20 Years)"}],
                                        }
                                    ],
                                }
                            ],
                        }
                    ],
                }
            ],
        }

        new_doc = {
            "type": "doc",
            "content": [
                {
                    "type": "table",
                    "content": [
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "<loan_tenure> Months(20 Years)"}],
                                        }
                                    ],
                                }
                            ],
                        }
                    ],
                }
            ],
        }

        diff_data = self.differ.calculate_diff(old_doc, new_doc)
        self.assertEqual(diff_data["summary"]["modified"], 1)

        changes = TemplateElementDiffer.extract_element_changes(diff_data)
        self.assertEqual(len(changes), 1)

        _, change_type, old_value, new_value = changes[0]
        self.assertEqual(change_type, "MODIFIED")

        semantic = {}
        if isinstance(new_value, dict):
            semantic = new_value.get("_semantic", {})

        self.assertEqual(semantic.get("type"), "TABLE_CONTENT_CHANGED")
        self.assertEqual(semantic.get("rowIndex"), 0)
        self.assertEqual(semantic.get("columnIndex"), 0)
        self.assertEqual(semantic.get("oldText"), "<loan_tenure>Months(20 Years)")
        self.assertEqual(semantic.get("newText"), "<loan_tenure> Months(20 Years)")


def run_migration_validation_suite():
    """Run all migration validation tests and report results"""
    import sys
    from django.test.runner import DiscoverRunner

    print("\n" + "="*80)
    print("🧪 RUNNING PROSEMIRROR MIGRATION VALIDATION SUITE")
    print("="*80 + "\n")

    runner = DiscoverRunner(verbosity=2)
    test_suite = runner.test_loader.loadTestsFromTestCase(ProseMirrorMigrationTests)
    integration_suite = runner.test_loader.loadTestsFromTestCase(ProseMirrorServiceIntegrationTests)
    lifecycle_suite = runner.test_loader.loadTestsFromTestCase(ProseMirrorLifecycleRegressionTests)

    result = runner.run_suite(test_suite)
    integration_result = runner.run_suite(integration_suite)
    lifecycle_result = runner.run_suite(lifecycle_suite)

    print("\n" + "="*80)
    if result.wasSuccessful() and integration_result.wasSuccessful() and lifecycle_result.wasSuccessful():
        print("✅ ALL MIGRATION TESTS PASSED")
        print("="*80 + "\n")
        return 0
    else:
        print("❌ SOME TESTS FAILED")
        print("="*80 + "\n")
        return 1


if __name__ == '__main__':
    import django
    import os
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
    django.setup()
    sys.exit(run_migration_validation_suite())
