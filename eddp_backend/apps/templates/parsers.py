"""
Enterprise DOCX parser for template import.

Canonical import target is ProseMirror JSON. The parser also emits enterprise
metadata used by layout validation, semantic diagnostics, and variable audits.
"""

from __future__ import annotations

import base64
import re
from collections import Counter
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.document import Document as WordDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

VARIABLE_PATTERN = re.compile(
    r"\{\{\s*([A-Za-z][A-Za-z0-9_]*)\s*\}\}|"
    r"\{\s*([A-Za-z][A-Za-z0-9_]*)\s*\}|"
    r"<\s*([A-Za-z][A-Za-z0-9_]*)\s*>"
)
HEADING_STYLE_PATTERN = re.compile(r"heading\s*([1-6])", flags=re.IGNORECASE)

DYNAMIC_TABLE_VARIABLES = {
    "ADDRESS_TABLE",
    "SIGNATURE_TABLE",
    "AMORTIZATION_TABLE",
    "PAYMENT_SCHEDULE",
    "CO_APPLICANT_TABLE",
}
SIGNATURE_VARIABLES = {
    "SIGNATURE",
    "AUTHORIZED_SIGNATORY",
    "CO_APPLICANT_SIGNATURE",
    "SIGNATURE_TABLE",
}
IMAGE_VARIABLES = {
    "CUSTOMER_PHOTO",
    "PROPERTY_IMAGE",
    "QR_CODE",
    "COMPANY_LOGO",
}

VARIABLE_NODE_BY_CATEGORY = {
    "SIMPLE": "variableChip",
    "DYNAMIC_TABLE": "dynamicTableVariable",
    "SIGNATURE": "signatureVariable",
    "IMAGE": "imagePlaceholderVariable",
}

VARIABLE_INLINE_NODE_TYPES = set(VARIABLE_NODE_BY_CATEGORY.values())

HIGHLIGHT_COLOR_MAP = {
    "AUTO": "#ffff00",
    "BLACK": "#000000",
    "BLUE": "#0000ff",
    "BRIGHT_GREEN": "#00ff00",
    "DARK_BLUE": "#00008b",
    "DARK_RED": "#8b0000",
    "DARK_YELLOW": "#b8860b",
    "GRAY_25": "#d9d9d9",
    "GRAY_50": "#808080",
    "GREEN": "#008000",
    "PINK": "#ffc0cb",
    "RED": "#ff0000",
    "TEAL": "#008080",
    "TURQUOISE": "#40e0d0",
    "VIOLET": "#8a2be2",
    "WHITE": "#ffffff",
    "YELLOW": "#ffff00",
}

TEXTUAL_SECTION_HEADINGS = {
    "customer details": "CUSTOMER_INFORMATION_BLOCK",
    "customer information": "CUSTOMER_INFORMATION_BLOCK",
    "loan details": "LOAN_DETAILS_SECTION",
    "loan information": "LOAN_DETAILS_SECTION",
    "address": "ADDRESS_SECTION",
    "address details": "ADDRESS_SECTION",
    "authorised signatory": "SIGNATURE_SECTION",
    "authorized signatory": "SIGNATURE_SECTION",
    "signature": "SIGNATURE_SECTION",
}

EMU_PER_PIXEL = 9525.0


class EnterpriseVariableRecognitionEngine:
    """Detects placeholders and normalizes them to enterprise variable chips."""

    def __init__(self) -> None:
        self._token_occurrences: Counter[str] = Counter()
        self._category_occurrences: Counter[str] = Counter()
        self._format_occurrences: Counter[str] = Counter()
        self._token_formats: Dict[str, set[str]] = {}

    @staticmethod
    def normalize_token(raw_token: str) -> str:
        token = re.sub(r"[^A-Za-z0-9_]", "_", raw_token.strip().upper())
        token = re.sub(r"_+", "_", token).strip("_")
        return token

    def _categorize(self, token: str) -> str:
        if token in IMAGE_VARIABLES or token.endswith("_IMAGE") or token.endswith("_PHOTO"):
            return "IMAGE"

        if token in DYNAMIC_TABLE_VARIABLES or token.endswith("_TABLE") or token.endswith("_SCHEDULE"):
            return "DYNAMIC_TABLE"

        if token in SIGNATURE_VARIABLES or "SIGNATURE" in token:
            return "SIGNATURE"

        return "SIMPLE"

    def tokenize(self, text: str) -> List[Dict[str, Any]]:
        if not text:
            return []

        parts: List[Dict[str, Any]] = []
        cursor = 0

        for match in VARIABLE_PATTERN.finditer(text):
            start, end = match.span()
            if start > cursor:
                parts.append({"type": "text", "value": text[cursor:start]})

            raw_name = match.group(1) or match.group(2) or match.group(3) or ""
            token = self.normalize_token(raw_name)
            if not token:
                parts.append({"type": "text", "value": match.group(0)})
                cursor = end
                continue

            matched_value = match.group(0)
            if matched_value.startswith("{{"):
                source_format = "DOUBLE_BRACE"
            elif matched_value.startswith("{"):
                source_format = "BRACE"
            else:
                source_format = "ANGLE"

            category = self._categorize(token)
            self._token_occurrences[token] += 1
            self._category_occurrences[category] += 1
            self._format_occurrences[source_format] += 1
            self._token_formats.setdefault(token, set()).add(source_format)

            parts.append(
                {
                    "type": "variable",
                    "token": token,
                    "source_format": source_format,
                    "category": category,
                }
            )
            cursor = end

        if cursor < len(text):
            parts.append({"type": "text", "value": text[cursor:]})

        return parts

    def summary(self) -> Dict[str, Any]:
        variables: List[Dict[str, Any]] = []
        for token in sorted(self._token_occurrences.keys()):
            category = self._categorize(token)
            variables.append(
                {
                    "token": token,
                    "normalized": f"<{token}>",
                    "category": category,
                    "occurrences": self._token_occurrences[token],
                    "source_formats": sorted(self._token_formats.get(token, set())),
                }
            )

        return {
            "total_detected": int(sum(self._token_occurrences.values())),
            "normalized_format": "<VARIABLE_NAME>",
            "by_category": {
                "simple": int(self._category_occurrences.get("SIMPLE", 0)),
                "dynamic_table": int(self._category_occurrences.get("DYNAMIC_TABLE", 0)),
                "signature": int(self._category_occurrences.get("SIGNATURE", 0)),
                "image": int(self._category_occurrences.get("IMAGE", 0)),
            },
            "by_source_format": {
                "brace": int(self._format_occurrences.get("BRACE", 0)),
                "angle": int(self._format_occurrences.get("ANGLE", 0)),
                "double_brace": int(self._format_occurrences.get("DOUBLE_BRACE", 0)),
            },
            "variables": variables,
        }


class NumberingResolver:
    """Resolves Word numbering metadata to ordered or bullet list semantics."""

    def __init__(self, document: WordDocument) -> None:
        self._num_to_abstract: Dict[int, int] = {}
        self._abstract_level_format: Dict[int, Dict[int, str]] = {}
        self._load_numbering(document)

    @staticmethod
    def _local_name(tag: str) -> str:
        return tag.split("}")[-1] if isinstance(tag, str) else ""

    @staticmethod
    def _attr_by_local_name(element: Any, attr_name: str) -> Optional[str]:
        for key, value in getattr(element, "attrib", {}).items():
            local = key.split("}")[-1]
            if local == attr_name:
                return str(value)
        return None

    def _load_numbering(self, document: WordDocument) -> None:
        try:
            numbering = document.part.numbering_part.numbering_definitions._numbering
        except Exception:
            return

        for abstract in numbering.xpath(".//*[local-name()='abstractNum']"):
            abstract_id_raw = self._attr_by_local_name(abstract, "abstractNumId")
            if abstract_id_raw is None:
                continue

            try:
                abstract_id = int(abstract_id_raw)
            except (TypeError, ValueError):
                continue

            level_mapping: Dict[int, str] = {}
            for lvl in abstract.xpath(".//*[local-name()='lvl']"):
                ilvl_raw = self._attr_by_local_name(lvl, "ilvl") or "0"
                try:
                    ilvl = int(ilvl_raw)
                except (TypeError, ValueError):
                    ilvl = 0

                num_fmt: Optional[str] = None
                for num_fmt_node in lvl.xpath(".//*[local-name()='numFmt']"):
                    num_fmt_raw = self._attr_by_local_name(num_fmt_node, "val")
                    if num_fmt_raw:
                        num_fmt = str(num_fmt_raw).lower()
                        break

                if num_fmt:
                    level_mapping[ilvl] = num_fmt

            self._abstract_level_format[abstract_id] = level_mapping

        for num in numbering.xpath(".//*[local-name()='num']"):
            num_id_raw = self._attr_by_local_name(num, "numId")
            if num_id_raw is None:
                continue

            abstract_id_raw: Optional[str] = None
            for abstract_ref in num.xpath(".//*[local-name()='abstractNumId']"):
                abstract_id_raw = self._attr_by_local_name(abstract_ref, "val")
                if abstract_id_raw is not None:
                    break

            if abstract_id_raw is None:
                continue

            try:
                num_id = int(num_id_raw)
                abstract_id = int(abstract_id_raw)
            except (TypeError, ValueError):
                continue

            self._num_to_abstract[num_id] = abstract_id

    def resolve_list_type(self, num_id: int, level: int, style_name: str, text: str) -> str:
        abstract_id = self._num_to_abstract.get(num_id)
        if abstract_id is not None:
            level_formats = self._abstract_level_format.get(abstract_id, {})
            num_fmt = level_formats.get(level) or level_formats.get(0)
            if num_fmt == "bullet":
                return "bulletList"
            if num_fmt:
                return "orderedList"

        style_lower = style_name.lower()
        if "bullet" in style_lower:
            return "bulletList"

        if "number" in style_lower or re.match(r"\s*\d+[.)]\s+", text or ""):
            return "orderedList"

        return "orderedList"


class ProseMirrorDocumentParser:
    """Parse DOCX files into enterprise ProseMirror JSON payloads."""

    def __init__(self) -> None:
        self.variable_engine = EnterpriseVariableRecognitionEngine()
        self.numbering_resolver: Optional[NumberingResolver] = None
        self.unsupported_features: List[Dict[str, Any]] = []
        self.layout_differences: List[Dict[str, Any]] = []
        self.formatting_differences: List[Dict[str, Any]] = []
        self.semantic_sections: List[Dict[str, Any]] = []
        self._unsupported_keys: set[str] = set()

    def parse(self, file_path_or_stream: Any) -> Dict[str, Any]:
        document = Document(file_path_or_stream)

        self.variable_engine = EnterpriseVariableRecognitionEngine()
        self.numbering_resolver = NumberingResolver(document)
        self.unsupported_features = []
        self.layout_differences = []
        self.formatting_differences = []
        self.semantic_sections = []
        self._unsupported_keys = set()

        self._inspect_document_level_features(document)

        body_content = self._parse_block_container(document, scope="body")
        if not body_content:
            body_content = [{"type": "paragraph"}]

        layout_metadata = self._build_layout_metadata(document, body_content)
        prosemirror_json = {
            "type": "doc",
            "content": body_content,
        }

        validation_report = self._build_validation_report(prosemirror_json, layout_metadata)

        return {
            "prosemirror_json": prosemirror_json,
            "layout": layout_metadata,
            "variable_summary": self.variable_engine.summary(),
            "semantic_sections": self.semantic_sections,
            "validation_report": validation_report,
        }

    def _inspect_document_level_features(self, document: WordDocument) -> None:
        root = document.part.element

        if root.xpath(".//*[local-name()='footnoteReference']"):
            self._add_unsupported_feature(
                code="FOOTNOTES",
                detail="Footnotes are not rendered as native editor annotations.",
                impact="medium",
                location="document",
            )

        if root.xpath(".//*[local-name()='endnoteReference']"):
            self._add_unsupported_feature(
                code="ENDNOTES",
                detail="Endnotes are not rendered as native editor annotations.",
                impact="medium",
                location="document",
            )

        if root.xpath(".//*[local-name()='commentRangeStart']"):
            self._add_unsupported_feature(
                code="WORD_COMMENTS",
                detail="Word comments are not imported into inline review comments.",
                impact="low",
                location="document",
            )

        if root.xpath(".//*[local-name()='txbxContent']"):
            self._add_unsupported_feature(
                code="TEXT_BOX",
                detail="Text boxes are retained as metadata but not rendered as floating shapes.",
                impact="high",
                location="document",
            )

        if self._contains_watermark(root):
            self._add_unsupported_feature(
                code="WATERMARK",
                detail="Watermarks are detected but not rendered as independent floating overlays.",
                impact="medium",
                location="document",
            )

    @staticmethod
    def _contains_watermark(root: Any) -> bool:
        for candidate in root.xpath(".//*[local-name()='shape' or local-name()='sp' or local-name()='docPr']"):
            for value in getattr(candidate, "attrib", {}).values():
                if "watermark" in str(value).lower():
                    return True
        return False

    def _parse_block_container(self, parent: Any, scope: str) -> List[Dict[str, Any]]:
        content: List[Dict[str, Any]] = []
        list_stack: List[Dict[str, Any]] = []

        for block in self._iter_blocks(parent):
            if isinstance(block, Paragraph):
                paragraph_blocks, list_info = self._parse_paragraph(block, scope)
                if not paragraph_blocks:
                    continue

                if list_info and len(paragraph_blocks) == 1 and paragraph_blocks[0].get("type") == "paragraph":
                    self._append_list_paragraph(
                        root_content=content,
                        list_stack=list_stack,
                        paragraph_node=paragraph_blocks[0],
                        level=int(list_info.get("level", 0)),
                        list_type=str(list_info.get("type", "bulletList")),
                    )
                    continue

                list_stack.clear()
                content.extend(paragraph_blocks)
                continue

            if isinstance(block, Table):
                list_stack.clear()
                content.append(self._parse_table(block, scope))

        list_stack.clear()
        return content

    def _iter_blocks(self, parent: Any) -> Iterable[Any]:
        if isinstance(parent, WordDocument):
            parent_elm = parent.element.body
            wrapper_parent = parent
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
            wrapper_parent = parent
        elif hasattr(parent, "_element"):
            parent_elm = parent._element
            wrapper_parent = parent
        else:
            return []

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, wrapper_parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, wrapper_parent)

    @staticmethod
    def _paragraph_has_substantive_content(paragraph: Paragraph) -> bool:
        text = (paragraph.text or "").strip()
        if text:
            return True

        for run in paragraph.runs:
            if run._element.xpath(".//*[local-name()='drawing']"):
                return True
            if run._element.xpath(".//*[local-name()='br']"):
                return True

        return False

    def _parse_paragraph(self, paragraph: Paragraph, scope: str) -> Tuple[List[Dict[str, Any]], Optional[Dict[str, Any]]]:
        paragraph_format = self._extract_paragraph_format(paragraph)
        list_info = self._resolve_list_info(paragraph)

        chunks: List[Dict[str, Any]] = []
        for run, hyperlink in self._iter_runs_with_links(paragraph):
            chunks.extend(self._parse_run_chunks(run, hyperlink))

        if not chunks and paragraph.text:
            chunks.extend(self._text_to_chunks(paragraph.text, marks=[]))

        if not chunks:
            if self._paragraph_has_substantive_content(paragraph):
                node = self._build_paragraph_node(
                    inline_nodes=[],
                    paragraph=paragraph,
                    paragraph_format=paragraph_format,
                    plain_text="",
                    node_type="paragraph",
                    heading_level=None,
                )
                self._register_semantic_role(node, paragraph, scope)
                return [node], list_info
            if self._should_preserve_empty_paragraph(scope, paragraph_format, list_info):
                node = self._build_paragraph_node(
                    inline_nodes=[],
                    paragraph=paragraph,
                    paragraph_format=paragraph_format,
                    plain_text="",
                    node_type="paragraph",
                    heading_level=None,
                )
                self._register_semantic_role(node, paragraph, scope)
                return [node], list_info
            return [], list_info

        node_type, heading_level = self._resolve_paragraph_node_type(paragraph)

        block_nodes: List[Dict[str, Any]] = []
        inline_buffer: List[Dict[str, Any]] = []
        first_textual_block = True

        for chunk in chunks:
            kind = chunk.get("kind")
            if kind == "inline":
                node = chunk.get("node")
                if isinstance(node, dict):
                    self._append_inline_node(inline_buffer, node)
                continue

            if inline_buffer:
                paragraph_node = self._build_paragraph_node(
                    inline_nodes=inline_buffer,
                    paragraph=paragraph,
                    paragraph_format=paragraph_format,
                    plain_text=self._inline_plain_text(inline_buffer),
                    node_type=node_type if first_textual_block else "paragraph",
                    heading_level=heading_level if first_textual_block else None,
                )
                self._register_semantic_role(paragraph_node, paragraph, scope)
                block_nodes.append(paragraph_node)
                inline_buffer = []
                first_textual_block = False

            block_node = chunk.get("node")
            if isinstance(block_node, dict):
                block_nodes.append(block_node)

        if inline_buffer or not block_nodes:
            paragraph_node = self._build_paragraph_node(
                inline_nodes=inline_buffer,
                paragraph=paragraph,
                paragraph_format=paragraph_format,
                plain_text=self._inline_plain_text(inline_buffer),
                node_type=node_type if first_textual_block else "paragraph",
                heading_level=heading_level if first_textual_block else None,
            )
            self._register_semantic_role(paragraph_node, paragraph, scope)
            block_nodes.append(paragraph_node)

        return block_nodes, list_info

    @staticmethod
    def _should_preserve_empty_paragraph(
        scope: str,
        paragraph_format: Dict[str, Any],
        list_info: Optional[Dict[str, Any]],
    ) -> bool:
        if list_info is not None:
            return True

        if scope == "body" or "table_cell" in scope:
            return True

        for key in ("lineSpacing", "spaceBefore", "spaceAfter", "firstLineIndent", "leftIndent", "rightIndent"):
            value = paragraph_format.get(key)
            if value not in (None, 0, 0.0):
                return True

        return False

    def _resolve_paragraph_node_type(self, paragraph: Paragraph) -> Tuple[str, Optional[int]]:
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        style_match = HEADING_STYLE_PATTERN.search(style_name)
        if style_match:
            try:
                return "heading", max(1, min(6, int(style_match.group(1))))
            except (TypeError, ValueError):
                return "heading", 1

        normalized_text = self._normalize_text(paragraph.text)
        if normalized_text and self._looks_like_title(paragraph, normalized_text):
            return "heading", 1

        if normalized_text and self._looks_like_section_heading(paragraph, normalized_text):
            return "heading", 2

        return "paragraph", None

    @staticmethod
    def _normalize_text(text: Optional[str]) -> str:
        value = (text or "").strip()
        return re.sub(r"\s+", " ", value)

    def _looks_like_title(self, paragraph: Paragraph, normalized_text: str) -> bool:
        if len(normalized_text) > 120:
            return False

        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            return False

        has_bold_run = any(bool(run.bold) for run in paragraph.runs if (run.text or "").strip())
        if not has_bold_run:
            return False

        alphabetic = re.sub(r"[^A-Za-z]", "", normalized_text)
        if not alphabetic:
            return False

        uppercase_ratio = sum(1 for char in alphabetic if char.isupper()) / max(1, len(alphabetic))
        return uppercase_ratio >= 0.6

    def _looks_like_section_heading(self, paragraph: Paragraph, normalized_text: str) -> bool:
        lowered = normalized_text.lower()
        if lowered in TEXTUAL_SECTION_HEADINGS:
            return True

        if len(normalized_text) > 80:
            return False

        has_bold_run = any(bool(run.bold) for run in paragraph.runs if (run.text or "").strip())
        return has_bold_run and lowered.endswith("details")

    def _extract_paragraph_format(self, paragraph: Paragraph) -> Dict[str, Any]:
        pf = paragraph.paragraph_format

        attrs: Dict[str, Any] = {}

        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            attrs["textAlign"] = "center"
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            attrs["textAlign"] = "right"
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            attrs["textAlign"] = "justify"
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            attrs["textAlign"] = "left"

        attrs["lineSpacing"] = self._safe_float(self._length_to_points(pf.line_spacing))
        attrs["spaceBefore"] = self._safe_float(self._length_to_points(pf.space_before))
        attrs["spaceAfter"] = self._safe_float(self._length_to_points(pf.space_after))
        attrs["firstLineIndent"] = self._safe_float(self._length_to_points(pf.first_line_indent))
        attrs["leftIndent"] = self._safe_float(self._length_to_points(pf.left_indent))
        attrs["rightIndent"] = self._safe_float(self._length_to_points(pf.right_indent))

        tab_stops: List[Dict[str, Any]] = []
        try:
            for stop in pf.tab_stops:
                tab_stops.append(
                    {
                        "position": self._safe_float(self._length_to_points(stop.position)),
                        "alignment": str(getattr(stop.alignment, "name", "LEFT")).upper(),
                    }
                )
        except Exception:
            tab_stops = []

        if tab_stops:
            attrs["tabStops"] = tab_stops

        return attrs

    @staticmethod
    def _safe_float(value: Optional[float]) -> Optional[float]:
        if value is None:
            return None
        try:
            return round(float(value), 2)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _length_to_points(value: Any) -> Optional[float]:
        if value is None:
            return None

        if hasattr(value, "pt"):
            try:
                return float(value.pt)
            except (TypeError, ValueError):
                return None

        if isinstance(value, (int, float)):
            return float(value)

        return None

    def _build_paragraph_node(
        self,
        inline_nodes: List[Dict[str, Any]],
        paragraph: Paragraph,
        paragraph_format: Dict[str, Any],
        plain_text: str,
        node_type: str,
        heading_level: Optional[int],
    ) -> Dict[str, Any]:
        attrs: Dict[str, Any] = {}

        if node_type == "heading":
            attrs["level"] = heading_level or 1

        text_align = paragraph_format.get("textAlign")
        if text_align:
            attrs["textAlign"] = text_align

        line_spacing = paragraph_format.get("lineSpacing")
        if line_spacing is not None:
            attrs["lineHeight"] = line_spacing

        space_before = paragraph_format.get("spaceBefore")
        if space_before is not None:
            attrs["spacingBefore"] = space_before

        space_after = paragraph_format.get("spaceAfter")
        if space_after is not None:
            attrs["spacingAfter"] = space_after

        first_line_indent = paragraph_format.get("firstLineIndent")
        if first_line_indent is not None:
            attrs["firstLineIndent"] = first_line_indent

        left_indent = paragraph_format.get("leftIndent")
        if left_indent is not None:
            attrs["leftIndent"] = left_indent

        right_indent = paragraph_format.get("rightIndent")
        if right_indent is not None:
            attrs["rightIndent"] = right_indent

        docx_attrs = {
            key: value
            for key, value in paragraph_format.items()
            if key not in {"textAlign"} and value is not None
        }
        if docx_attrs:
            attrs["docx"] = docx_attrs

        node: Dict[str, Any] = {"type": node_type}
        if attrs:
            node["attrs"] = attrs
        if inline_nodes:
            node["content"] = inline_nodes

        # Preserve role hints for downstream diagnostics.
        semantic_role = self._infer_semantic_role(plain_text)
        if semantic_role:
            node.setdefault("attrs", {})["semanticRole"] = semantic_role

        return node

    def _infer_semantic_role(self, plain_text: str) -> Optional[str]:
        normalized = self._normalize_text(plain_text)
        if not normalized:
            return None

        lowered = normalized.lower()
        if lowered in TEXTUAL_SECTION_HEADINGS:
            return TEXTUAL_SECTION_HEADINGS[lowered]

        if "authorised signatory" in lowered or "authorized signatory" in lowered:
            return "SIGNATURE_SECTION"

        if "signature" in lowered and len(lowered) <= 80:
            return "SIGNATURE_SECTION"

        if "address" in lowered and len(lowered) <= 80:
            return "ADDRESS_SECTION"

        if "loan" in lowered and "detail" in lowered:
            return "LOAN_DETAILS_SECTION"

        if "customer" in lowered and ("detail" in lowered or "information" in lowered):
            return "CUSTOMER_INFORMATION_BLOCK"

        return None

    def _register_semantic_role(self, node: Dict[str, Any], paragraph: Paragraph, scope: str) -> None:
        if scope != "body":
            return

        plain_text = self._extract_node_plain_text(node)
        role = self._infer_semantic_role(plain_text)
        if not role:
            return

        self.semantic_sections.append(
            {
                "role": role,
                "text": plain_text,
                "style": paragraph.style.name if paragraph.style is not None else "",
                "node_type": node.get("type"),
            }
        )

    @staticmethod
    def _extract_node_plain_text(node: Dict[str, Any]) -> str:
        content = node.get("content")
        if not isinstance(content, list):
            return ""

        parts: List[str] = []
        for child in content:
            if not isinstance(child, dict):
                continue

            node_type = child.get("type")
            if node_type == "text":
                parts.append(str(child.get("text") or ""))
            elif node_type == "hardBreak":
                parts.append("\n")
            elif node_type in VARIABLE_INLINE_NODE_TYPES:
                attrs = child.get("attrs") if isinstance(child.get("attrs"), dict) else {}
                token = str(attrs.get("field") or attrs.get("label") or "VARIABLE")
                parts.append(f"<{token}>")

        return re.sub(r"\s+", " ", "".join(parts)).strip()

    def _resolve_list_info(self, paragraph: Paragraph) -> Optional[Dict[str, Any]]:
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        style_lower = style_name.lower()

        num_pr = None
        if paragraph._p.pPr is not None:
            num_pr = paragraph._p.pPr.numPr

        if num_pr is not None and num_pr.numId is not None:
            try:
                num_id = int(num_pr.numId.val)
            except (TypeError, ValueError):
                num_id = 0

            try:
                level = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
            except (TypeError, ValueError):
                level = 0

            list_type = "orderedList"
            if self.numbering_resolver is not None:
                list_type = self.numbering_resolver.resolve_list_type(
                    num_id=num_id,
                    level=level,
                    style_name=style_name,
                    text=paragraph.text,
                )

            return {"level": max(0, level), "type": list_type}

        style_level_match = re.search(r"(\d+)$", style_lower)
        style_level = int(style_level_match.group(1)) - 1 if style_level_match else 0

        if "list bullet" in style_lower:
            return {"level": max(0, style_level), "type": "bulletList"}

        if "list number" in style_lower:
            return {"level": max(0, style_level), "type": "orderedList"}

        return None

    def _append_list_paragraph(
        self,
        root_content: List[Dict[str, Any]],
        list_stack: List[Dict[str, Any]],
        paragraph_node: Dict[str, Any],
        level: int,
        list_type: str,
    ) -> None:
        if level < 0:
            level = 0

        if list_stack and level > list_stack[-1]["level"] + 1:
            level = list_stack[-1]["level"] + 1

        while list_stack and list_stack[-1]["level"] > level:
            list_stack.pop()

        while list_stack and list_stack[-1]["level"] == level and list_stack[-1]["type"] != list_type:
            list_stack.pop()

        if not list_stack:
            list_node = {"type": list_type, "content": []}
            root_content.append(list_node)
            list_stack.append({"level": 0, "type": list_type, "node": list_node, "last_item": None})

        while list_stack and list_stack[-1]["level"] < level:
            parent_entry = list_stack[-1]
            parent_item = parent_entry.get("last_item")
            if parent_item is None:
                parent_item = {"type": "listItem", "content": [{"type": "paragraph"}]}
                parent_entry["node"]["content"].append(parent_item)
                parent_entry["last_item"] = parent_item

            nested_list = {"type": list_type, "content": []}
            parent_item.setdefault("content", []).append(nested_list)
            list_stack.append(
                {
                    "level": parent_entry["level"] + 1,
                    "type": list_type,
                    "node": nested_list,
                    "last_item": None,
                }
            )

        target = list_stack[-1]
        list_item = {"type": "listItem", "content": [paragraph_node]}
        target["node"]["content"].append(list_item)
        target["last_item"] = list_item

    def _iter_runs_with_links(self, paragraph: Paragraph) -> Iterable[Tuple[Run, Optional[str]]]:
        for child in paragraph._p.iterchildren():
            tag = self._local_name(child.tag)
            if tag == "r":
                yield Run(child, paragraph), None
                continue

            if tag != "hyperlink":
                continue

            hyperlink = self._resolve_hyperlink(paragraph, child)
            for run_node in child.iterchildren():
                if self._local_name(run_node.tag) != "r":
                    continue
                yield Run(run_node, paragraph), hyperlink

    @staticmethod
    def _local_name(tag: str) -> str:
        return tag.split("}")[-1] if isinstance(tag, str) else ""

    def _resolve_hyperlink(self, paragraph: Paragraph, hyperlink_node: Any) -> Optional[str]:
        rel_id = hyperlink_node.get(qn("r:id"))
        anchor = hyperlink_node.get(qn("w:anchor"))

        href: Optional[str] = None
        if rel_id:
            try:
                relationship = paragraph.part.rels.get(rel_id)
                if relationship is not None:
                    href = str(getattr(relationship, "target_ref", "") or "")
            except Exception:
                href = None

        if anchor:
            if href:
                return f"{href}#{anchor}"
            return f"#{anchor}"

        return href

    def _parse_run_chunks(self, run: Run, hyperlink: Optional[str]) -> List[Dict[str, Any]]:
        marks = self._build_run_marks(run, hyperlink)
        chunks: List[Dict[str, Any]] = []
        saw_text_node = False

        for child in run._element.iterchildren():
            child_tag = self._local_name(child.tag)

            if child_tag == "t":
                saw_text_node = True
                chunks.extend(self._text_to_chunks(child.text or "", marks))
                continue

            if child_tag == "tab":
                chunks.extend(self._text_to_chunks("\t", marks))
                continue

            if child_tag == "instrText":
                instr_text = (child.text or "").strip().upper()
                if "PAGE" in instr_text:
                    chunks.extend(self._text_to_chunks("<PAGE_NUMBER>", marks=[]))
                continue

            if child_tag in {"br", "cr"}:
                break_type = self._attr_by_local_name(child, "type")
                if break_type == "page":
                    chunks.append(
                        {
                            "kind": "block",
                            "node": {
                                "type": "horizontalRule",
                                "attrs": {"docxPageBreak": True},
                            },
                        }
                    )
                else:
                    chunks.append({"kind": "inline", "node": {"type": "hardBreak"}})
                continue

            if child_tag == "drawing":
                image_node = self._extract_image_node(child, run.part)
                if image_node is not None:
                    chunks.append({"kind": "block", "node": image_node})
                continue

            if child_tag == "object":
                self._add_unsupported_feature(
                    code="OLE_OBJECT",
                    detail="OLE embedded objects are retained as metadata only.",
                    impact="high",
                    location="run",
                )

        if not saw_text_node and run.text:
            chunks.extend(self._text_to_chunks(run.text, marks))

        return chunks

    def _build_run_marks(self, run: Run, hyperlink: Optional[str]) -> List[Dict[str, Any]]:
        marks: List[Dict[str, Any]] = []
        text_style_attrs: Dict[str, Any] = {}

        if run.bold:
            marks.append({"type": "bold"})

        if run.italic:
            marks.append({"type": "italic"})

        if run.underline:
            marks.append({"type": "underline"})

        if bool(run.font.strike) or bool(run.font.double_strike):
            marks.append({"type": "strike"})

        if run.font.name:
            text_style_attrs["fontFamily"] = str(run.font.name)

        if run.font.size is not None and hasattr(run.font.size, "pt"):
            try:
                text_style_attrs["fontSize"] = f"{round(float(run.font.size.pt), 2)}pt"
            except (TypeError, ValueError):
                pass

        if run.font.color is not None and run.font.color.rgb is not None:
            text_style_attrs["color"] = f"#{str(run.font.color.rgb)}".lower()

        font_spacing = getattr(run.font, "spacing", None)
        if font_spacing is not None and hasattr(font_spacing, "pt"):
            try:
                text_style_attrs["letterSpacing"] = f"{round(float(font_spacing.pt), 2)}pt"
            except (TypeError, ValueError):
                pass

        if run.font.all_caps:
            text_style_attrs["textTransform"] = "uppercase"

        if run.font.small_caps:
            text_style_attrs["fontVariant"] = "small-caps"

        if run.font.superscript:
            text_style_attrs["verticalAlign"] = "super"
        elif run.font.subscript:
            text_style_attrs["verticalAlign"] = "sub"

        highlight_name = getattr(run.font.highlight_color, "name", None)
        if highlight_name:
            marks.append(
                {
                    "type": "highlight",
                    "attrs": {
                        "color": HIGHLIGHT_COLOR_MAP.get(str(highlight_name).upper(), "#ffff00"),
                    },
                }
            )

        shading_fill = self._extract_run_shading_fill(run)
        if shading_fill:
            text_style_attrs["backgroundColor"] = shading_fill

        if text_style_attrs:
            marks.append({"type": "textStyle", "attrs": text_style_attrs})

        if hyperlink:
            marks.append({"type": "link", "attrs": {"href": hyperlink}})

        return marks

    def _extract_run_shading_fill(self, run: Run) -> Optional[str]:
        rpr = run._element.rPr
        if rpr is None:
            return None

        for shd in rpr.xpath(".//*[local-name()='shd']"):
            fill = self._attr_by_local_name(shd, "fill")
            if fill and fill.lower() not in {"auto", "000000"}:
                return f"#{fill.lower()}"

        return None

    def _text_to_chunks(self, text: str, marks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        chunks: List[Dict[str, Any]] = []
        for segment in self.variable_engine.tokenize(text):
            segment_type = segment.get("type")
            if segment_type == "text":
                value = str(segment.get("value") or "")
                if not value:
                    continue
                text_node: Dict[str, Any] = {"type": "text", "text": value}
                if marks:
                    text_node["marks"] = marks
                chunks.append({"kind": "inline", "node": text_node})
                continue

            if segment_type == "variable":
                token = str(segment.get("token") or "").strip()
                if not token:
                    continue
                category = str(segment.get("category") or "SIMPLE")
                source_format = str(segment.get("source_format") or "")
                node_type = VARIABLE_NODE_BY_CATEGORY.get(category, "variableChip")
                chunks.append(
                    {
                        "kind": "inline",
                        "node": {
                            "type": node_type,
                            "attrs": {
                                "field": token,
                                "label": token,
                                "category": category,
                                "normalized": f"<{token}>",
                                "source_format": source_format,
                            },
                        },
                    }
                )

        return chunks

    def _extract_image_node(self, drawing_element: Any, part: Any) -> Optional[Dict[str, Any]]:
        blips = drawing_element.xpath(".//*[local-name()='blip']")
        if not blips:
            if drawing_element.xpath(".//*[local-name()='chart' or local-name()='diagram']"):
                self._add_unsupported_feature(
                    code="SMART_ART_OR_CHART",
                    detail="Charts and SmartArt are detected but not rendered as editable chart nodes.",
                    impact="high",
                    location="drawing",
                )
            return None

        for blip in blips:
            embed_id = self._attr_by_local_name(blip, "embed")
            if not embed_id:
                continue

            image_part = getattr(part, "related_parts", {}).get(embed_id)
            if image_part is None:
                continue

            blob = getattr(image_part, "blob", None)
            content_type = str(getattr(image_part, "content_type", "") or "")
            if not blob or not content_type.startswith("image/"):
                continue

            encoded = base64.b64encode(blob).decode("ascii")
            src = f"data:{content_type};base64,{encoded}"

            attrs: Dict[str, Any] = {"src": src}
            alt_text = self._extract_drawing_alt_text(drawing_element)
            if alt_text:
                attrs["alt"] = alt_text

            dimensions = self._extract_drawing_dimensions(drawing_element)
            if dimensions.get("width") is not None:
                attrs["width"] = dimensions["width"]
            if dimensions.get("height") is not None:
                attrs["height"] = dimensions["height"]

            return {"type": "image", "attrs": attrs}

        return None

    def _extract_drawing_alt_text(self, drawing_element: Any) -> Optional[str]:
        for doc_pr in drawing_element.xpath(".//*[local-name()='docPr']"):
            descr = self._attr_by_local_name(doc_pr, "descr")
            if descr:
                return descr
            title = self._attr_by_local_name(doc_pr, "title")
            if title:
                return title
            name = self._attr_by_local_name(doc_pr, "name")
            if name:
                return name
        return None

    def _extract_drawing_dimensions(self, drawing_element: Any) -> Dict[str, Optional[int]]:
        for extent in drawing_element.xpath(".//*[local-name()='extent']"):
            cx_raw = self._attr_by_local_name(extent, "cx")
            cy_raw = self._attr_by_local_name(extent, "cy")

            width = self._emu_to_pixels(cx_raw)
            height = self._emu_to_pixels(cy_raw)
            if width is not None or height is not None:
                return {"width": width, "height": height}

        return {"width": None, "height": None}

    @staticmethod
    def _emu_to_pixels(value: Optional[str]) -> Optional[int]:
        if value is None:
            return None
        try:
            return max(1, int(round(float(value) / EMU_PER_PIXEL)))
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _attr_by_local_name(element: Any, attr_name: str) -> Optional[str]:
        for key, value in getattr(element, "attrib", {}).items():
            local = key.split("}")[-1]
            if local == attr_name:
                return str(value)
        return None

    def _append_inline_node(self, inline_buffer: List[Dict[str, Any]], node: Dict[str, Any]) -> None:
        if node.get("type") != "text":
            inline_buffer.append(node)
            return

        text_value = str(node.get("text") or "")
        if not text_value:
            return

        marks = node.get("marks")
        if (
            inline_buffer
            and inline_buffer[-1].get("type") == "text"
            and inline_buffer[-1].get("marks") == marks
        ):
            inline_buffer[-1]["text"] = f"{inline_buffer[-1].get('text', '')}{text_value}"
            return

        inline_buffer.append(node)

    def _inline_plain_text(self, inline_nodes: List[Dict[str, Any]]) -> str:
        parts: List[str] = []
        for node in inline_nodes:
            node_type = node.get("type")
            if node_type == "text":
                parts.append(str(node.get("text") or ""))
            elif node_type == "hardBreak":
                parts.append("\n")
            elif node_type in VARIABLE_INLINE_NODE_TYPES:
                attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
                token = str(attrs.get("field") or attrs.get("label") or "VARIABLE")
                parts.append(f"<{token}>")
        return "".join(parts)

    def _parse_table(self, table: Table, scope: str) -> Dict[str, Any]:
        rows: List[Dict[str, Any]] = []

        for row_index, row in enumerate(table.rows):
            row_cells: List[Dict[str, Any]] = []
            header_row = self._is_header_row(row, row_index)

            for cell in row.cells:
                if self._is_vertical_merge_continuation(cell):
                    continue

                cell_node_type = "tableHeader" if header_row else "tableCell"
                cell_node: Dict[str, Any] = {"type": cell_node_type}

                cell_attrs = self._extract_cell_attrs(cell)
                if cell_attrs:
                    cell_node["attrs"] = cell_attrs

                cell_content = self._parse_block_container(cell, scope=f"{scope}_table_cell")
                if not cell_content:
                    cell_content = [{"type": "paragraph"}]

                cell_node["content"] = cell_content
                row_cells.append(cell_node)

            if row_cells:
                row_node: Dict[str, Any] = {"type": "tableRow", "content": row_cells}
                row_attrs = self._extract_row_attrs(row)
                if row_attrs:
                    row_node["attrs"] = row_attrs
                rows.append(row_node)

        if not rows:
            rows = [
                {
                    "type": "tableRow",
                    "content": [{"type": "tableCell", "content": [{"type": "paragraph"}]}],
                }
            ]

        attrs: Dict[str, Any] = self._extract_table_attrs(table, scope)
        if table.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            attrs["textAlign"] = "center"
        elif table.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            attrs["textAlign"] = "right"
        elif table.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            attrs["textAlign"] = "left"

        node: Dict[str, Any] = {
            "type": "table",
            "content": rows,
        }
        if attrs:
            node["attrs"] = attrs

        return node

    @staticmethod
    def _is_header_row(row: _Row, row_index: int) -> bool:
        if row_index != 0:
            return False

        bold_runs = 0
        total_runs = 0
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if not (run.text or "").strip():
                        continue
                    total_runs += 1
                    if run.bold:
                        bold_runs += 1

        if total_runs == 0:
            return False

        return (bold_runs / total_runs) >= 0.5

    @staticmethod
    def _is_vertical_merge_continuation(cell: _Cell) -> bool:
        tc_pr = cell._tc.tcPr
        if tc_pr is None or tc_pr.vMerge is None:
            return False

        merge_val = tc_pr.vMerge.val
        if merge_val is None:
            return True

        return str(merge_val).lower() == "continue"

    def _extract_cell_attrs(self, cell: _Cell) -> Dict[str, Any]:
        attrs: Dict[str, Any] = {}

        tc_pr = cell._tc.tcPr
        if tc_pr is not None and tc_pr.gridSpan is not None:
            try:
                colspan = int(tc_pr.gridSpan.val)
                if colspan > 1:
                    attrs["colspan"] = colspan
            except (TypeError, ValueError):
                pass

        if cell.width is not None and hasattr(cell.width, "pt"):
            try:
                width_px = int(round((float(cell.width.pt) / 72.0) * 96.0))
                if width_px > 0:
                    attrs["colwidth"] = [width_px]
            except (TypeError, ValueError):
                pass

        vertical_alignment = getattr(cell, "vertical_alignment", None)
        if vertical_alignment is not None:
            attrs["docxVerticalAlign"] = str(getattr(vertical_alignment, "name", vertical_alignment)).upper()

        cell_fill = self._extract_cell_fill(cell)
        if cell_fill:
            attrs["backgroundColor"] = cell_fill

        borders = self._extract_cell_borders(cell)
        if borders:
            attrs["docxBorders"] = borders

        return attrs

    def _extract_table_attrs(self, table: Table, scope: str) -> Dict[str, Any]:
        attrs: Dict[str, Any] = {}

        if "table_cell" in scope:
            attrs["docxNested"] = True

        if table.style is not None and getattr(table.style, "name", None):
            attrs["docxTableStyle"] = str(table.style.name)

        try:
            width = table._tbl.tblPr.tblW if table._tbl.tblPr is not None else None
            width_raw = self._attr_by_local_name(width, "w") if width is not None else None
            width_type = self._attr_by_local_name(width, "type") if width is not None else None
            if width_raw:
                attrs["docxTableWidth"] = width_raw
            if width_type:
                attrs["docxTableWidthType"] = width_type
        except Exception:
            pass

        return attrs

    def _extract_row_attrs(self, row: _Row) -> Dict[str, Any]:
        attrs: Dict[str, Any] = {}

        tr_pr = row._tr.trPr
        if tr_pr is None:
            return attrs

        heights = tr_pr.xpath(".//*[local-name()='trHeight']")
        if not heights:
            return attrs

        height_raw = self._attr_by_local_name(heights[0], "val")
        if not height_raw:
            return attrs

        try:
            # DOCX row height val is twips.
            attrs["docxHeightPt"] = round(float(height_raw) / 20.0, 2)
        except (TypeError, ValueError):
            return {}

        return attrs

    def _extract_cell_fill(self, cell: _Cell) -> Optional[str]:
        tc_pr = cell._tc.tcPr
        if tc_pr is None:
            return None

        shd = tc_pr.xpath(".//*[local-name()='shd']")
        if not shd:
            return None

        fill = self._attr_by_local_name(shd[0], "fill")
        if not fill or fill.lower() in {"auto", "000000"}:
            return None

        return f"#{fill.lower()}"

    def _extract_cell_borders(self, cell: _Cell) -> Dict[str, Any]:
        tc_pr = cell._tc.tcPr
        if tc_pr is None:
            return {}

        border_nodes = tc_pr.xpath(".//*[local-name()='tcBorders']")
        if not border_nodes:
            return {}

        borders: Dict[str, Any] = {}
        valid_edges = {"top", "bottom", "left", "right", "insideH", "insideV"}
        for border in border_nodes[0].iterchildren():
            edge = self._local_name(border.tag)
            if edge not in valid_edges:
                continue

            edge_payload: Dict[str, Any] = {}
            for attr in ("val", "sz", "color", "space"):
                attr_value = self._attr_by_local_name(border, attr)
                if attr_value:
                    edge_payload[attr] = attr_value

            if edge_payload:
                borders[edge] = edge_payload

        return borders

    def _build_layout_metadata(self, document: WordDocument, body_content: List[Dict[str, Any]]) -> Dict[str, Any]:
        sections: List[Dict[str, Any]] = []
        headers: List[Dict[str, Any]] = []
        footers: List[Dict[str, Any]] = []

        for index, section in enumerate(document.sections):
            section_data = self._extract_section_layout(section, index)
            sections.append(section_data)

            header_content = self._trim_empty_paragraphs(
                self._parse_block_container(section.header, scope=f"header_{index}")
            )
            footer_content = self._trim_empty_paragraphs(
                self._parse_block_container(section.footer, scope=f"footer_{index}")
            )

            if header_content:
                headers.append({"section_index": index, "content": header_content})
            if footer_content:
                footers.append({"section_index": index, "content": footer_content})

        first_section = sections[0] if sections else {}

        title = (document.core_properties.title or "").strip()
        if not title:
            title = self._derive_title_from_content(body_content)

        bookmarks = self._extract_bookmarks(document)

        return {
            "title": title,
            "page_size": first_section.get("page_size", "CUSTOM"),
            "page_orientation": first_section.get("orientation", "PORTRAIT"),
            "margins": first_section.get("margins", {}),
            "sections": sections,
            "headers": headers,
            "footers": footers,
            "bookmarks": bookmarks,
            "page_number_fields": self._count_page_number_fields(document),
        }

    def _extract_section_layout(self, section: Any, section_index: int) -> Dict[str, Any]:
        orientation_value = getattr(section, "orientation", WD_ORIENT.PORTRAIT)
        if orientation_value == WD_ORIENT.LANDSCAPE:
            orientation = "LANDSCAPE"
        else:
            orientation = "PORTRAIT"

        width_pt = self._safe_float(self._length_to_points(section.page_width))
        height_pt = self._safe_float(self._length_to_points(section.page_height))
        page_size = self._classify_page_size(width_pt, height_pt)

        section_data = {
            "section_index": section_index,
            "orientation": orientation,
            "page_size": page_size,
            "page_width_pt": width_pt,
            "page_height_pt": height_pt,
            "margins": {
                "top": self._safe_float(self._length_to_points(section.top_margin)),
                "bottom": self._safe_float(self._length_to_points(section.bottom_margin)),
                "left": self._safe_float(self._length_to_points(section.left_margin)),
                "right": self._safe_float(self._length_to_points(section.right_margin)),
            },
            "columns": self._extract_section_column_count(section),
        }

        if section_data["columns"] > 1:
            self._add_unsupported_feature(
                code="MULTI_COLUMN_LAYOUT",
                detail="Multi-column section layout is preserved as metadata but rendered as single flow.",
                impact="high",
                location=f"section:{section_index}",
            )
            self.layout_differences.append(
                {
                    "type": "MULTI_COLUMN_LAYOUT",
                    "section": section_index,
                    "message": "Editor currently renders multi-column section as single-column flow.",
                }
            )

        return section_data

    def _extract_section_column_count(self, section: Any) -> int:
        cols_nodes = section._sectPr.xpath(".//*[local-name()='cols']")
        if not cols_nodes:
            return 1

        col_count_raw = self._attr_by_local_name(cols_nodes[0], "num")
        if not col_count_raw:
            return 1

        try:
            return max(1, int(col_count_raw))
        except (TypeError, ValueError):
            return 1

    @staticmethod
    def _classify_page_size(width_pt: Optional[float], height_pt: Optional[float]) -> str:
        if width_pt is None or height_pt is None:
            return "CUSTOM"

        width_in = width_pt / 72.0
        height_in = height_pt / 72.0
        short_side = min(width_in, height_in)
        long_side = max(width_in, height_in)

        sizes = {
            "A4": (8.27, 11.69),
            "A3": (11.69, 16.54),
            "LETTER": (8.5, 11.0),
            "LEGAL": (8.5, 14.0),
        }

        for label, (w, h) in sizes.items():
            if abs(short_side - w) <= 0.2 and abs(long_side - h) <= 0.2:
                return label

        return "CUSTOM"

    def _derive_title_from_content(self, body_content: List[Dict[str, Any]]) -> str:
        for node in body_content:
            if not isinstance(node, dict):
                continue
            if node.get("type") == "heading":
                candidate = self._extract_node_plain_text(node)
                if candidate:
                    return candidate

        for node in body_content:
            if not isinstance(node, dict):
                continue
            if node.get("type") == "paragraph":
                candidate = self._extract_node_plain_text(node)
                if candidate:
                    return candidate[:120]

        return ""

    def _extract_bookmarks(self, document: WordDocument) -> List[Dict[str, Any]]:
        bookmarks: List[Dict[str, Any]] = []
        for bookmark in document.part.element.xpath(".//*[local-name()='bookmarkStart']"):
            name = self._attr_by_local_name(bookmark, "name")
            bookmark_id = self._attr_by_local_name(bookmark, "id")
            if not name or name.startswith("_"):
                continue
            bookmarks.append({"name": name, "id": bookmark_id or ""})

        return bookmarks

    def _count_page_number_fields(self, document: WordDocument) -> int:
        count = 0
        for section in document.sections:
            for paragraph in section.footer.paragraphs:
                if paragraph._p.xpath(".//*[local-name()='instrText' and contains(translate(text(), 'page', 'PAGE'), 'PAGE')]"):
                    count += 1
        return count

    @staticmethod
    def _trim_empty_paragraphs(content: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        trimmed = []
        for node in content:
            if not isinstance(node, dict):
                continue
            if node.get("type") != "paragraph":
                trimmed.append(node)
                continue
            child_content = node.get("content")
            if isinstance(child_content, list) and child_content:
                trimmed.append(node)
                continue

            attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
            docx_attrs = attrs.get("docx") if isinstance(attrs.get("docx"), dict) else {}
            preserve_due_to_spacing = any(
                docx_attrs.get(key) not in (None, 0, 0.0)
                for key in ("lineSpacing", "spaceBefore", "spaceAfter", "firstLineIndent", "leftIndent", "rightIndent")
            )
            if preserve_due_to_spacing:
                trimmed.append(node)
        return trimmed

    def _build_validation_report(
        self,
        prosemirror_json: Dict[str, Any],
        layout_metadata: Dict[str, Any],
    ) -> Dict[str, Any]:
        content = prosemirror_json.get("content") if isinstance(prosemirror_json.get("content"), list) else []
        node_counts = self._count_node_types(content)

        unsupported_count = len(self.unsupported_features)
        formatting_diff_count = len(self.formatting_differences)
        layout_diff_count = len(self.layout_differences)

        penalty = (unsupported_count * 0.02) + (formatting_diff_count * 0.005) + (layout_diff_count * 0.01)
        estimated_similarity = max(0.75, round(0.99 - penalty, 4))

        return {
            "target_similarity_range": "0.95-0.99",
            "estimated_similarity_score": estimated_similarity,
            "meets_target": estimated_similarity >= 0.95,
            "node_counts": node_counts,
            "formatting_differences": self.formatting_differences,
            "layout_differences": self.layout_differences,
            "unsupported_features": self.unsupported_features,
            "variable_detection": self.variable_engine.summary(),
            "layout_snapshot": {
                "page_size": layout_metadata.get("page_size"),
                "page_orientation": layout_metadata.get("page_orientation"),
                "section_count": len(layout_metadata.get("sections", [])),
                "header_count": len(layout_metadata.get("headers", [])),
                "footer_count": len(layout_metadata.get("footers", [])),
            },
        }

    def _count_node_types(self, nodes: List[Dict[str, Any]]) -> Dict[str, int]:
        counter: Counter[str] = Counter()

        def walk(node: Any) -> None:
            if not isinstance(node, dict):
                return
            node_type = str(node.get("type") or "unknown")
            counter[node_type] += 1
            content = node.get("content")
            if isinstance(content, list):
                for child in content:
                    walk(child)

        for node in nodes:
            walk(node)

        return {key: int(counter[key]) for key in sorted(counter.keys())}

    def _add_unsupported_feature(
        self,
        *,
        code: str,
        detail: str,
        impact: str,
        location: str,
    ) -> None:
        key = f"{code}|{detail}|{location}"
        if key in self._unsupported_keys:
            return

        self._unsupported_keys.add(key)
        self.unsupported_features.append(
            {
                "code": code,
                "detail": detail,
                "impact": impact,
                "location": location,
            }
        )
