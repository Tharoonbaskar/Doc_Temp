from __future__ import annotations

from io import BytesIO
from time import perf_counter
from typing import Any
import html
import re

from django.utils import timezone
from docx import Document as DocxDocument

from apps.common.choices import OutputFormatChoices
from apps.common.exceptions import ExternalServiceException, ValidationException

from ..models import GenerationRequest
from ..repositories import RuntimeEngineRepository
from .file_storage import FileStorageService


class DOCXGeneratorService:
    def __init__(
        self,
        repository: RuntimeEngineRepository | None = None,
        file_storage_service: FileStorageService | None = None,
    ) -> None:
        self.repository = repository or RuntimeEngineRepository()
        self.file_storage_service = file_storage_service or FileStorageService()

    @staticmethod
    def _log(
        logs: list[dict[str, Any]],
        *,
        stage: str,
        message: str,
        metadata: dict[str, Any] | None = None,
    ) -> None:
        logs.append(
            {
                "timestamp": timezone.now().isoformat(),
                "stage": stage,
                "message": message,
                "metadata": metadata or {},
            }
        )

    @staticmethod
    def _normalize_file_name(file_name: str, generation_request: GenerationRequest) -> str:
        candidate = (file_name or "").strip()
        if not candidate:
            candidate = f"{generation_request.code}-{timezone.now().strftime('%Y%m%d%H%M%S')}.docx"
        if not candidate.lower().endswith(".docx"):
            candidate = f"{candidate}.docx"
        return candidate

    @staticmethod
    def _html_to_lines(value: str) -> list[str]:
        if not value:
            return []

        cleaned = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", value)
        cleaned = re.sub(r"(?i)<br\s*/?>", "\n", cleaned)
        cleaned = re.sub(r"(?i)</(p|div|h1|h2|h3|h4|h5|h6|li|tr|table|section|article|header|footer|main)>", "\n", cleaned)
        cleaned = re.sub(r"(?i)<(td|th)[^>]*>", " ", cleaned)
        cleaned = re.sub(r"<[^>]+>", "", cleaned)

        text = html.unescape(cleaned)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return lines

    def generate_docx(
        self,
        *,
        generation_request: GenerationRequest,
        html_content: str,
        header_html: str = "",
        footer_html: str = "",
        file_name: str = "",
        document_title: str = "",
    ) -> dict[str, Any]:
        if generation_request is None:
            raise ValidationException(detail="generation_request is required for DOCX generation.")
        if not isinstance(html_content, str) or not html_content.strip():
            raise ValidationException(detail="html_content is required for DOCX generation.")

        execution_log: list[dict[str, Any]] = []
        started_at = perf_counter()
        target_file_name = self._normalize_file_name(file_name, generation_request)

        self._log(
            execution_log,
            stage="DOCX_GENERATION_START",
            message="DOCX generation started.",
            metadata={"file_name": target_file_name},
        )

        try:
            document = DocxDocument()

            title = (document_title or "").strip()
            if title:
                document.add_heading(title, level=1)

            header_lines = self._html_to_lines(header_html)
            footer_lines = self._html_to_lines(footer_html)
            body_lines = self._html_to_lines(html_content)

            if header_lines:
                section_header = document.sections[0].header
                section_header.is_linked_to_previous = False
                paragraph = section_header.paragraphs[0] if section_header.paragraphs else section_header.add_paragraph()
                paragraph.text = " | ".join(header_lines)

            if footer_lines:
                section_footer = document.sections[0].footer
                section_footer.is_linked_to_previous = False
                paragraph = section_footer.paragraphs[0] if section_footer.paragraphs else section_footer.add_paragraph()
                paragraph.text = " | ".join(footer_lines)

            if not body_lines:
                body_lines = [""]

            for line in body_lines:
                document.add_paragraph(line)

            output = BytesIO()
            document.save(output)
            docx_bytes = output.getvalue()
        except Exception as exc:
            raise ExternalServiceException(detail=f"DOCX generation failed: {exc}") from exc

        storage_result = self.file_storage_service.save_bytes(
            content=docx_bytes,
            file_name=target_file_name,
            subdirectory="generated-documents/docx",
        )

        generated_document = self.repository.upsert_generated_document(
            generation_request=generation_request,
            file_name=storage_result["file_name"],
            file_path=storage_result["relative_path"],
            file_type=OutputFormatChoices.DOCX,
            file_size=storage_result["file_size"],
            checksum=storage_result["checksum"],
        )

        duration_ms = round((perf_counter() - started_at) * 1000, 3)
        self._log(
            execution_log,
            stage="DOCX_GENERATION_COMPLETE",
            message="DOCX generation completed.",
            metadata={
                "file_size": storage_result["file_size"],
                "duration_ms": duration_ms,
                "storage_backend": storage_result["storage_backend"],
            },
        )

        return {
            "generated_document_id": str(generated_document.id),
            "generated_document_code": generated_document.code,
            "file_name": generated_document.file_name,
            "file_path": generated_document.file_path,
            "file_type": generated_document.file_type,
            "file_size": generated_document.file_size,
            "checksum": generated_document.checksum,
            "storage_backend": storage_result["storage_backend"],
            "duration_ms": duration_ms,
            "execution_log": execution_log,
        }
